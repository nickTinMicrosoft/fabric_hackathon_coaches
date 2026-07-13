#!/usr/bin/env python3
"""Generate a Word document that documents the University database."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).parent / "University_Database.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)

# --------------------------------------------------------------------------
# Content model: tables -> columns (name, type, key, description)
# --------------------------------------------------------------------------
TABLES = [
    ("Department", "Dimension",
     "Academic departments that own programs, courses and professors.",
     [
        ("DeptID", "INT", "PK", "Unique department identifier."),
        ("DeptCode", "VARCHAR(8)", "", "Short code (e.g. CS, MATH, ENG, BUS, SCI)."),
        ("DeptName", "VARCHAR(100)", "", "Full department name."),
     ]),
    ("Term", "Dimension",
     "Academic terms in chronological order; the time axis for all facts.",
     [
        ("TermID", "INT", "PK", "Unique term identifier."),
        ("TermName", "VARCHAR(40)", "", "Display name (e.g. Fall 2024)."),
        ("StartDate", "DATE", "", "First day of the term."),
        ("EndDate", "DATE", "", "Last day of the term."),
        ("TermOrder", "INT", "", "Chronological sequence (1 = earliest)."),
     ]),
    ("Building", "Dimension",
     "Physical buildings on campus.",
     [
        ("BldID", "INT", "PK", "Unique building identifier."),
        ("BldName", "VARCHAR(100)", "", "Building name."),
     ]),
    ("Classroom", "Dimension",
     "Rooms where sections are taught, with type and seating capacity.",
     [
        ("RoomID", "INT", "PK", "Unique room identifier."),
        ("RoomName", "VARCHAR(40)", "", "Room label (e.g. SH-Lab-A)."),
        ("BldID", "INT", "FK", "Building the room belongs to (Building.BldID)."),
        ("RoomType", "VARCHAR(20)", "", "Lecture, Lab, Research or Seminar."),
        ("Capacity", "INT", "", "Maximum seats."),
     ]),
    ("Program", "Dimension",
     "Degree programs a student can be enrolled in, with credit requirement.",
     [
        ("ProgramID", "INT", "PK", "Unique program identifier."),
        ("ProgramName", "VARCHAR(100)", "", "Program title (e.g. BS Computer Science)."),
        ("DeptID", "INT", "FK", "Owning department (Department.DeptID)."),
        ("DegreeType", "VARCHAR(20)", "", "Degree awarded (BA or BS)."),
        ("RequiredCredits", "INT", "", "Credits required to graduate."),
     ]),
    ("Professor", "Dimension",
     "Instructors who teach sections; carries rank and salary for cost analysis.",
     [
        ("ProfID", "INT", "PK", "Unique professor identifier."),
        ("ProfName", "VARCHAR(100)", "", "Professor name."),
        ("DeptID", "INT", "FK", "Home department (Department.DeptID)."),
        ("Rank", "VARCHAR(30)", "", "Adjunct, Assistant, Associate or Full."),
        ("AnnualSalary", "INT", "", "Annual salary; basis for section instructor cost."),
     ]),
    ("Course", "Dimension",
     "Catalog of courses offered by departments.",
     [
        ("CrsID", "INT", "PK", "Unique course identifier."),
        ("CrsNmbr", "VARCHAR(12)", "", "Course number (e.g. CS101)."),
        ("CrsName", "VARCHAR(120)", "", "Course title."),
        ("DeptID", "INT", "FK", "Owning department (Department.DeptID)."),
        ("Credits", "INT", "", "Credit hours awarded."),
        ("CrsLevel", "INT", "", "Course level (100, 200, ...)."),
     ]),
    ("Student", "Dimension",
     "The student population with the demographic attributes used to explain and "
     "predict outcomes. StudentID is the stable key shared with the Cosmos "
     "graduate-exam documents.",
     [
        ("StudentID", "INT", "PK", "Unique student identifier (shared with Cosmos)."),
        ("StudentName", "VARCHAR(100)", "", "Student full name."),
        ("ProgramID", "INT", "FK", "Program the student pursues (Program.ProgramID)."),
        ("EnrollmentTermID", "INT", "FK", "Term the student first enrolled (Term.TermID)."),
        ("ExpectedGradTermID", "INT", "FK", "Planned graduation term (Term.TermID)."),
        ("FirstGen", "BIT", "", "1 = first-generation college student."),
        ("Residency", "VARCHAR(12)", "", "InState or OutOfState."),
        ("AdmissionScore", "INT", "", "Admission strength (900-1600)."),
        ("FinancialAidTier", "INT", "", "Financial need: 0 (none) to 3 (high)."),
        ("CurrentStatus", "VARCHAR(12)", "", "Active, Withdrawn or Graduated."),
     ]),
    ("CourseSection", "Fact",
     "A specific offering of a course in a term, taught by a professor in a room. "
     "Carries capacity, seats filled and cost so cost-effectiveness is measurable.",
     [
        ("SectionID", "INT", "PK", "Unique section identifier."),
        ("CrsID", "INT", "FK", "Course being offered (Course.CrsID)."),
        ("TermID", "INT", "FK", "Term of the offering (Term.TermID)."),
        ("ProfID", "INT", "FK", "Instructor (Professor.ProfID)."),
        ("RoomID", "INT", "FK", "Room used (Classroom.RoomID)."),
        ("Capacity", "INT", "", "Seat capacity for the section."),
        ("SeatsFilled", "INT", "", "Number of students enrolled."),
        ("InstructorCost", "DECIMAL(10,2)", "", "Instructor cost to run the section."),
        ("RoomCost", "DECIMAL(10,2)", "", "Facility cost to run the section."),
     ]),
    ("Enrollment", "Fact",
     "One row per student per section: the grade outcome (as quality points) and "
     "whether the student withdrew.",
     [
        ("EnrollmentID", "INT", "PK", "Unique enrollment identifier."),
        ("StudentID", "INT", "FK", "Enrolled student (Student.StudentID)."),
        ("SectionID", "INT", "FK", "Section taken (CourseSection.SectionID)."),
        ("CrsID", "INT", "FK", "Course taken (Course.CrsID)."),
        ("TermID", "INT", "FK", "Term of enrollment (Term.TermID)."),
        ("LetterGrade", "VARCHAR(2)", "", "Letter grade; NULL if withdrawn."),
        ("GradePoints", "DECIMAL(3,2)", "", "Quality points 0.00-4.00; NULL if withdrawn."),
        ("Withdrawn", "BIT", "", "1 = student withdrew from the section."),
        ("AttemptNumber", "INT", "", "Attempt number for repeated courses."),
     ]),
    ("StudentTermStatus", "Fact",
     "One row per student per term - the retention and progress backbone. Combines "
     "the term outcome with engagement signals that act as early warning indicators.",
     [
        ("StatusID", "INT", "PK", "Unique row identifier."),
        ("StudentID", "INT", "FK", "Student (Student.StudentID)."),
        ("TermID", "INT", "FK", "Term (Term.TermID)."),
        ("Status", "VARCHAR(12)", "", "Enrolled, Withdrawn, Graduated or Leave."),
        ("TermGpa", "DECIMAL(3,2)", "", "GPA earned in this term."),
        ("CumGpa", "DECIMAL(3,2)", "", "Cumulative GPA through this term."),
        ("CreditsAttempted", "INT", "", "Credits attempted in the term."),
        ("CreditsEarned", "INT", "", "Credits passed in the term."),
        ("CumCreditsEarned", "INT", "", "Total credits earned to date (vs RequiredCredits)."),
        ("AttendancePct", "INT", "", "Attendance rate 0-100 (engagement signal)."),
        ("LmsLogins", "INT", "", "Learning-management-system logins (engagement signal)."),
        ("AdvisingVisits", "INT", "", "Advising appointments attended (engagement signal)."),
     ]),
    ("CourseReviewSignal", "Fact",
     "AI-distilled signals from the free-text course reviews stored in Blob Storage. "
     "The raw review text stays in the storage account at BlobPath; the structured "
     "sentiment, themes and recommendation produced by the language model are "
     "materialised here for analytics.",
     [
        ("ReviewID", "INT", "PK", "Unique review identifier."),
        ("StudentID", "INT", "FK", "Reviewing student (Student.StudentID)."),
        ("SectionID", "INT", "FK", "Section reviewed (CourseSection.SectionID)."),
        ("CrsID", "INT", "FK", "Course reviewed (Course.CrsID)."),
        ("TermID", "INT", "FK", "Term of the review (Term.TermID)."),
        ("SubmittedDate", "DATE", "", "Date the review was submitted."),
        ("SentimentScore", "DECIMAL(3,2)", "", "Sentiment from -1.00 (neg) to 1.00 (pos)."),
        ("ThemePacing", "BIT", "", "1 = review mentions course pacing."),
        ("ThemeWorkload", "BIT", "", "1 = review mentions workload."),
        ("ThemeInstructor", "BIT", "", "1 = review mentions the instructor."),
        ("ThemeMaterials", "BIT", "", "1 = review mentions course materials."),
        ("WouldRecommend", "BIT", "", "1 = student would recommend the course."),
        ("SummaryText", "VARCHAR(300)", "", "One-line summary of the review."),
        ("BlobPath", "VARCHAR(200)", "", "Path to the raw review text in Blob Storage."),
     ]),
]

RELATIONSHIPS = [
    ("Program", "Department", "each program belongs to one department"),
    ("Professor", "Department", "each professor has a home department"),
    ("Course", "Department", "each course is owned by a department"),
    ("Classroom", "Building", "each room is located in a building"),
    ("Student", "Program", "each student pursues one program"),
    ("Student", "Term", "enrollment and expected-graduation terms"),
    ("CourseSection", "Course / Term / Professor / Classroom", "a section ties them together"),
    ("Enrollment", "Student / CourseSection / Course / Term", "a student's record in a section"),
    ("StudentTermStatus", "Student / Term", "one status row per student per term"),
    ("CourseReviewSignal", "Student / CourseSection / Course / Term", "a review of a section"),
]

QUESTIONS = [
    ("How do we increase student retention?",
     "StudentTermStatus gives one row per student per term, so a student who does not "
     "appear in the next term has left. Joined with Student demographics and the "
     "engagement signals (AttendancePct, LmsLogins, AdvisingVisits), the leading "
     "indicators of drop-out become directly queryable."),
    ("How do we increase graduation rates?",
     "CumCreditsEarned in StudentTermStatus, compared against Program.RequiredCredits, "
     "shows each student's progress and pace toward the degree over time."),
    ("How do we make the business more cost-effective?",
     "CourseSection carries InstructorCost, RoomCost, Capacity and SeatsFilled. Combined "
     "with credits earned in Enrollment, this yields cost per completed credit and "
     "surfaces under-filled or high-cost sections."),
    ("How do we use AI to understand the student experience?",
     "CourseReviewSignal holds language-model output (sentiment, themes, recommendation) "
     "distilled from the free-text course reviews in Blob Storage, turning unstructured "
     "feedback into analytics-ready columns that explain the numbers."),
]


def shade_row(row, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


# --------------------------------------------------------------------------
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# Title page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("University Analytics Database")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Schema & Data Dictionary")
r.font.size = Pt(14)
r.font.color.rgb = GREY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("Azure SQL Database  \u2022  Student Demo Environment")
rm.font.size = Pt(10)
rm.font.color.rgb = GREY

doc.add_paragraph()

# Overview
heading(doc, "1. Overview", 1)
doc.add_paragraph(
    "This database models a university's students, courses and academic operations "
    "with a design centred on analytics. It follows a star-schema style: dimension "
    "tables describe the who and what (students, programs, courses, professors, "
    "terms and rooms), while fact tables capture outcomes and the leading indicators "
    "that explain them (term-by-term progress, enrollment grades, section cost and "
    "AI-distilled course-review sentiment)."
)
doc.add_paragraph(
    "The design deliberately supports four business questions: increasing student "
    "retention, increasing graduation rates, improving cost-effectiveness, and using "
    "AI to understand the student experience."
)
doc.add_paragraph(
    "StudentID is a stable identifier shared with the graduate-entrance-exam documents "
    "stored in Cosmos DB, so the structured SQL data and the semi-structured exam "
    "results can be joined per student."
)

# Table catalog
heading(doc, "2. Table Catalog", 1)
cat = doc.add_table(rows=1, cols=3)
cat.style = 'Light Grid Accent 1'
cat.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = cat.rows[0]
set_cell_text(hdr.cells[0], "Table", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
set_cell_text(hdr.cells[1], "Type", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
set_cell_text(hdr.cells[2], "Purpose", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
shade_row(hdr, "1F3864")
for name, kind, desc, _cols in TABLES:
    row = cat.add_row()
    set_cell_text(row.cells[0], name, bold=True)
    set_cell_text(row.cells[1], kind)
    set_cell_text(row.cells[2], desc)

doc.add_paragraph()

# Relationships
heading(doc, "3. Relationships", 1)
doc.add_paragraph(
    "Fact tables reference dimension tables through foreign keys. The key "
    "relationships are:"
)
for child, parent, note in RELATIONSHIPS:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f"{child} \u2192 {parent}: ")
    r.bold = True
    p.add_run(note + ".")

# Data dictionary
heading(doc, "4. Data Dictionary", 1)
for name, kind, desc, cols in TABLES:
    heading(doc, f"4.{[t[0] for t in TABLES].index(name)+1}  {name}", 2)
    dp = doc.add_paragraph()
    dr = dp.add_run(f"{kind} table. ")
    dr.italic = True
    dr.font.color.rgb = GREY
    dp.add_run(desc)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light List Accent 1'
    h = tbl.rows[0]
    set_cell_text(h.cells[0], "Column", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(h.cells[1], "Type", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(h.cells[2], "Key", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_text(h.cells[3], "Description", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    shade_row(h, "1F3864")
    for cname, ctype, key, cdesc in cols:
        r = tbl.add_row()
        set_cell_text(r.cells[0], cname, bold=(key == "PK"))
        set_cell_text(r.cells[1], ctype)
        set_cell_text(r.cells[2], key)
        set_cell_text(r.cells[3], cdesc)
    for r in tbl.rows:
        r.cells[0].width = Inches(1.6)
        r.cells[1].width = Inches(1.2)
        r.cells[2].width = Inches(0.5)
        r.cells[3].width = Inches(3.4)
    doc.add_paragraph()

# Analytics questions
heading(doc, "5. What the Data Answers", 1)
for i, (question, answer) in enumerate(QUESTIONS, start=1):
    hp = doc.add_paragraph()
    hr = hp.add_run(f"{i}. {question}")
    hr.bold = True
    hr.font.color.rgb = NAVY
    doc.add_paragraph(answer)

# Conventions
heading(doc, "6. Conventions & Notes", 1)
for note in [
    "Grades are stored as quality points (GradePoints, 0.00-4.00) so GPA can be "
    "aggregated directly without a lookup; LetterGrade is provided for readability.",
    "A withdrawn enrollment has NULL LetterGrade and GradePoints and Withdrawn = 1.",
    "Engagement signals (AttendancePct, LmsLogins, AdvisingVisits) live on "
    "StudentTermStatus and act as early-warning indicators for retention models.",
    "Course-review raw text is kept in Blob Storage; only the AI-distilled signals "
    "are stored in CourseReviewSignal, with BlobPath pointing back to the source.",
    "Credit hours are 3 per course; a student graduates when CumCreditsEarned "
    "reaches the program's RequiredCredits.",
    "All identifier columns are integers and all foreign keys are enforced to keep "
    "the model referentially consistent.",
]:
    doc.add_paragraph(note, style='List Bullet')

doc.save(OUT)
print(f"Wrote {OUT}")
